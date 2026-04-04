import docx
doc = docx.Document(r'c:\Users\minsoo\Desktop\capstone\농산물 가격 예측 프로젝트 계획서.docx')
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text)
print("=== TABLES ===")
for table in doc.tables:
    for row in table.rows:
        row_text = [cell.text.strip() for cell in row.cells]
        print(" | ".join(row_text))
